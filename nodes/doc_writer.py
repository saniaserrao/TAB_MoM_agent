from __future__ import annotations

import logging
import os
import re
import sys
import tempfile
from collections import defaultdict
from datetime import date
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "utils"))
from blob import upload_blob, make_session_blob_path

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------

_HEADING_COLOR = RGBColor(0x00, 0x47, 0xAB)
_BODY_FONT     = "Calibri"
_HEADING_FONT  = "Calibri Light"
_BASE_SIZE     = Pt(11)
_H1_SIZE       = Pt(14)
_H2_SIZE       = Pt(12)
_TITLE_SIZE    = Pt(18)

_ROLE_LABELS = {
    "ta":      "Microsoft Innovation Hub",
    "msft":    "Microsoft",
    "client":  "Client",
    "unknown": "Unknown",
}

_FLEX_LABELS = {
    "use_case_hypothesis":     "Use Case Hypotheses",
    "roi_signal":              "ROI Signals",
    "stakeholder_concern":     "Stakeholder Concerns",
    "technical_requirement":   "Technical Requirements",
    "azure_service_mentioned": "Azure Services Mentioned",
    "integration_point":       "Integration Points",
    "prototype_scope_item":    "Prototype Scope Items",
    "build_constraint":        "Build Constraints",
    "timeline_signal":         "Timeline Signals",
    "solution_component":      "Solution Components",
    "vendor_mention":          "Vendor Mentions",
    "risk_flag":               "Risk Flags",
}


# ---------------------------------------------------------------------------
# Designation lookup — built from participants list in state
# ---------------------------------------------------------------------------

def _build_designation_map(participants: list) -> dict[str, str]:
    """
    Build a name → designation lookup from the participants list supplied by the UI.
    Each participant: { "name": str, "designation": str, "role": str }
    Returns { "Eric Johnson": "CTO", "Eric": "CTO", ... } (full name + first name keys).
    """
    result: dict[str, str] = {}
    for p in (participants or []):
        name        = (p.get("name") or "").strip()
        designation = (p.get("designation") or "").strip()
        if name and designation:
            result[name] = designation
            first = name.split()[0]
            if first not in result:
                result[first] = designation
    return result


def _resolve_owner_label(owner_name: str, owner_role: str,
                         designation_map: dict[str, str]) -> str:
    """
    Return the best human-readable label for an action item owner.
    Priority: designation from participants > role label > raw role string.
    """
    if owner_name:
        desig = designation_map.get(owner_name)
        if not desig:
            first = owner_name.split()[0]
            desig = designation_map.get(first)
        if desig:
            return desig
    if owner_role and owner_role != "unknown":
        return _ROLE_LABELS.get(owner_role, owner_role.title())
    return ""


def _derive_filename(state: dict) -> str:
    """
    MoM_<SessionId>.docx
    Session ID is already canonical: <Company>_<EngagementType>_<DDMmmYYYY>
    """
    session_id = state.get("session_id", "session")
    safe_sid   = re.sub(r"[^\w\-]", "", session_id)[:60]
    return f"MoM_{safe_sid}.docx"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _h1(doc: Document, text: str) -> None:
    p   = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    run.font.name      = _HEADING_FONT
    run.font.size      = _H1_SIZE
    run.font.bold      = True
    run.font.color.rgb = _HEADING_COLOR


def _h2(doc: Document, text: str) -> None:
    p   = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    run.font.name      = _HEADING_FONT
    run.font.size      = _H2_SIZE
    run.font.bold      = True
    run.font.color.rgb = _HEADING_COLOR


def _body(doc: Document, text: str) -> None:
    p   = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.name = _BODY_FONT
    run.font.size = _BASE_SIZE


def _bullet(doc: Document, text: str, bold: bool = False) -> None:
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = _BODY_FONT
    run.font.size = _BASE_SIZE
    run.bold      = bold


def _meta_line(doc: Document, label: str, value: str) -> None:
    p  = doc.add_paragraph(style="Normal")
    r1 = p.add_run(label)
    r1.bold      = True
    r1.font.name = _BODY_FONT
    r1.font.size = _BASE_SIZE
    r2 = p.add_run(value)
    r2.font.name = _BODY_FONT
    r2.font.size = _BASE_SIZE


# ---------------------------------------------------------------------------
# Agenda-item grouping
# ---------------------------------------------------------------------------

def _keywords(text: str) -> set[str]:
    stopwords = {"this", "that", "with", "from", "they", "have", "will", "been",
                 "were", "their", "which", "about", "would", "could", "also",
                 "more", "than", "when", "what", "into", "some", "there", "then"}
    return {w for w in re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
            if w not in stopwords}


def _best_match(text: str, prob_keywords: list) -> int:
    kw     = _keywords(text)
    scores = [len(kw & pk) for pk in prob_keywords]
    best   = max(range(len(scores)), key=lambda i: scores[i])
    return best if scores[best] > 0 else -1


def _group_into_agenda_items(state: dict) -> list[dict]:
    problems  = state.get("problem_statements") or []
    actions   = state.get("action_items") or []
    queries   = state.get("client_queries") or []
    decisions = state.get("key_decisions") or []
    flex_tags = state.get("flexible_tags") or []
    summary   = state.get("summary") or ""

    # Separate discussion tags from other flex tags
    discussion_map: dict[str, str] = {}
    other_flex: list = []
    for tag in flex_tags:
        if tag.get("type") == "discussion":
            discussion_map[tag.get("speaker", "")] = tag.get("content", "")
        else:
            other_flex.append(tag)

    if not problems:
        return [{
            "title":      "General Discussion",
            "discussion": summary,
            "actions":    actions,
            "queries":    queries,
            "decisions":  decisions,
            "flex_tags":  other_flex,
        }]

    titles  = [p.get("content", "").strip() for p in problems]
    prob_kw = [_keywords(t) for t in titles]

    def _assign(items: list, key: str) -> tuple:
        buckets: list[list] = [[] for _ in problems]
        unmatched = []
        for item in items:
            text        = item.get(key) or item.get("description") or item.get("content") or ""
            agenda_topic = item.get("agenda_topic", "")
            if agenda_topic:
                for i, title in enumerate(titles):
                    if title.lower() == agenda_topic.lower():
                        buckets[i].append(item)
                        break
                else:
                    idx = _best_match(text, prob_kw)
                    (buckets[idx] if idx >= 0 else unmatched).append(item)
            else:
                idx = _best_match(text, prob_kw)
                (buckets[idx] if idx >= 0 else unmatched).append(item)
        return buckets, unmatched

    a_bkts, a_left = _assign(actions,    "description")
    q_bkts, q_left = _assign(queries,    "content")
    d_bkts, d_left = _assign(decisions,  "content")
    f_bkts, f_left = _assign(other_flex, "content")

    items = []
    for i, title in enumerate(titles):
        discussion = discussion_map.get(title, "") or (summary if i == 0 else "")
        items.append({
            "title":      title,
            "discussion": discussion,
            "actions":    a_bkts[i],
            "queries":    q_bkts[i],
            "decisions":  d_bkts[i],
            "flex_tags":  f_bkts[i],
        })

    if any([a_left, q_left, d_left, f_left]):
        items.append({
            "title":      "Additional Items",
            "discussion": "",
            "actions":    a_left,
            "queries":    q_left,
            "decisions":  d_left,
            "flex_tags":  f_left,
        })

    return items


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _derive_meeting_title(state: dict) -> str:
    """<Company> — <EngagementType> Meeting Minutes"""
    engagement_type = (state.get("engagement_type_detected")
                       or state.get("engagement_type") or "Meeting")
    company = state.get("company") or ""

    if not company:
        # Fall back to session_id first segment
        skip = {"smoke", "week", "bew", "ads", "poc", "test", "session",
                "mom", "demo", "unknown", "minimal"}
        for part in re.split(r"[-_]", state.get("session_id", "")):
            if len(part) > 2 and not part.isdigit() and part.lower() not in skip:
                company = part.capitalize()
                break

    if company:
        return f"{company} — {engagement_type} Meeting Minutes"
    return f"{engagement_type} Meeting Minutes"


def _set_para_spacing(p, space_before_pt: float = 0,
                      space_after_pt: float = 0) -> None:
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after  = Pt(space_after_pt)


def _build_cover(doc: Document, state: dict) -> None:
    today_str = date.today().strftime("%B %d, %Y")
    title_str = _derive_meeting_title(state)

    title_p   = doc.add_paragraph(style="Title")
    title_run = title_p.add_run(title_str)
    title_run.font.name      = _HEADING_FONT
    title_run.font.size      = _TITLE_SIZE
    title_run.font.color.rgb = _HEADING_COLOR
    title_p.alignment        = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(title_p, space_before_pt=0, space_after_pt=8)

    _meta_line(doc, "Date: ",       today_str)
    _meta_line(doc, "Session ID: ", state.get("session_id", "—"))

    speaker_map = state.get("speaker_map") or {}
    chairs      = [n for n, r in speaker_map.items() if r == "ta"]
    if chairs:
        _meta_line(doc, "Facilitator: ", chairs[0])

    grouped: dict[str, list[str]] = defaultdict(list)
    for name, role in speaker_map.items():
        grouped[role].append(name)

    # Build designation map for cover page attendee list
    participants    = state.get("participants") or []
    designation_map = _build_designation_map(participants)

    for role in ["ta", "msft", "client", "unknown"]:
        names = grouped.get(role)
        if names:
            p = doc.add_paragraph(style="Normal")
            _set_para_spacing(p, space_before_pt=0, space_after_pt=2)
            r1 = p.add_run(f"{_ROLE_LABELS.get(role, role.title())}: ")
            r1.bold      = True
            r1.font.name = _BODY_FONT
            r1.font.size = _BASE_SIZE
            # Include designation next to each name where available
            name_parts = []
            for n in sorted(names):
                desig = designation_map.get(n) or designation_map.get(n.split()[0])
                name_parts.append(f"{n} ({desig})" if desig else n)
            r2 = p.add_run(", ".join(name_parts))
            r2.font.name = _BODY_FONT
            r2.font.size = _BASE_SIZE

    sp = doc.add_paragraph()
    _set_para_spacing(sp, space_before_pt=4, space_after_pt=0)


def _build_agenda_items(doc: Document, state: dict,
                        designation_map: dict[str, str]) -> None:
    for idx, item in enumerate(_group_into_agenda_items(state), start=1):
        p_h1 = doc.add_paragraph(style="Heading 1")
        run  = p_h1.add_run(f"{idx}. {item['title']}")
        run.font.name      = _HEADING_FONT
        run.font.size      = _H1_SIZE
        run.font.bold      = True
        run.font.color.rgb = _HEADING_COLOR
        _set_para_spacing(p_h1, space_before_pt=14, space_after_pt=4)

        if item.get("discussion"):
            p_h2 = doc.add_paragraph(style="Heading 2")
            r = p_h2.add_run("Discussions")
            r.font.name = _HEADING_FONT; r.font.size = _H2_SIZE
            r.font.bold = True; r.font.color.rgb = _HEADING_COLOR
            _set_para_spacing(p_h2, space_before_pt=6, space_after_pt=2)
            p = doc.add_paragraph(style="Normal")
            p.add_run(item["discussion"]).font.name = _BODY_FONT
            _set_para_spacing(p, space_before_pt=0, space_after_pt=4)

        if item.get("decisions"):
            p_h2 = doc.add_paragraph(style="Heading 2")
            r = p_h2.add_run("Decisions")
            r.font.name = _HEADING_FONT; r.font.size = _H2_SIZE
            r.font.bold = True; r.font.color.rgb = _HEADING_COLOR
            _set_para_spacing(p_h2, space_before_pt=6, space_after_pt=2)
            for d in item["decisions"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(d.get("content", "")).font.name = _BODY_FONT
                _set_para_spacing(p, space_before_pt=0, space_after_pt=2)

        if item.get("actions"):
            p_h2 = doc.add_paragraph(style="Heading 2")
            r = p_h2.add_run("Action Items")
            r.font.name = _HEADING_FONT; r.font.size = _H2_SIZE
            r.font.bold = True; r.font.color.rgb = _HEADING_COLOR
            _set_para_spacing(p_h2, space_before_pt=6, space_after_pt=2)
            for ai in item["actions"]:
                desc     = ai.get("description", "")
                owner    = ai.get("owner_name", "")
                role     = ai.get("owner_role", "")
                deadline = ai.get("deadline")
                role_lbl = _resolve_owner_label(owner, role, designation_map)
                parts    = [desc]
                if owner:
                    parts.append(f"({owner}" + (f" — {role_lbl}" if role_lbl else "") + ")")
                if deadline:
                    parts.append(f"· Due: {deadline}")
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("  ".join(parts)).font.name = _BODY_FONT
                _set_para_spacing(p, space_before_pt=0, space_after_pt=2)

        if item.get("queries"):
            p_h2 = doc.add_paragraph(style="Heading 2")
            r = p_h2.add_run("Open Questions")
            r.font.name = _HEADING_FONT; r.font.size = _H2_SIZE
            r.font.bold = True; r.font.color.rgb = _HEADING_COLOR
            _set_para_spacing(p_h2, space_before_pt=6, space_after_pt=2)
            for q in item["queries"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(q.get("content", "")).font.name = _BODY_FONT
                _set_para_spacing(p, space_before_pt=0, space_after_pt=2)

        flex = item.get("flex_tags") or []
        if flex:
            p_h2 = doc.add_paragraph(style="Heading 2")
            r = p_h2.add_run("Key Insights")
            r.font.name = _HEADING_FONT; r.font.size = _H2_SIZE
            r.font.bold = True; r.font.color.rgb = _HEADING_COLOR
            _set_para_spacing(p_h2, space_before_pt=6, space_after_pt=2)
            by_type: dict[str, list] = defaultdict(list)
            for tag in flex:
                by_type[tag.get("type", "other")].append(tag)
            for tag_type, tags in by_type.items():
                display = _FLEX_LABELS.get(tag_type,
                                           tag_type.replace("_", " ").title())
                lbl = doc.add_paragraph(style="Normal")
                run = lbl.add_run(display)
                run.bold = True; run.font.name = _BODY_FONT
                run.font.size = _BASE_SIZE; run.font.color.rgb = _HEADING_COLOR
                _set_para_spacing(lbl, space_before_pt=4, space_after_pt=2)
                for tag in tags:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(tag.get("content", "")).font.name = _BODY_FONT
                    _set_para_spacing(p, space_before_pt=0, space_after_pt=2)


def _build_action_summary(doc: Document, state: dict,
                          designation_map: dict[str, str]) -> None:
    all_actions = state.get("action_items") or []
    if not all_actions:
        return

    _h1(doc, "Action Items Summary")

    col_widths = [Inches(3.2), Inches(1.6), Inches(1.5), Inches(1.2)]
    headers    = ["Description", "Owner", "Designation", "Deadline"]

    tbl = doc.add_table(rows=1 + len(all_actions), cols=4)
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = col_widths[i]
        _set_cell_bg(cell, "0047AB")
        run = cell.paragraphs[0].add_run(h)
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.name      = _BODY_FONT
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, ai in enumerate(all_actions, start=1):
        owner    = ai.get("owner_name", "")
        role     = ai.get("owner_role", "")
        role_lbl = _resolve_owner_label(owner, role, designation_map)
        values   = [
            ai.get("description", ""),
            owner,
            role_lbl or "—",
            ai.get("deadline") or "—",
        ]
        row = tbl.rows[row_idx]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            if row_idx % 2 == 0:
                _set_cell_bg(cell, "EEF3FB")
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = _BODY_FONT


# ---------------------------------------------------------------------------
# Node
# ---------------------------------------------------------------------------

def doc_writer_node(state: dict) -> dict:
    if state.get("error"):
        logger.warning("doc_writer_node: upstream error, passing through.")
        return state

    session_id   = state.get("session_id", "unknown_session")
    company      = (state.get("company") or "Client").strip().replace(" ", "_")
    filename     = _derive_filename(state)
    blob_path    = make_session_blob_path(company, session_id, filename)

    logger.info("doc_writer_node: building '%s'.", filename)

    doc = Document()
    doc.styles["Normal"].font.name = _BODY_FONT
    doc.styles["Normal"].font.size = _BASE_SIZE

    participants    = state.get("participants") or []
    designation_map = _build_designation_map(participants)
    logger.info("doc_writer_node: designation_map has %d entries.", len(designation_map))

    _build_cover(doc, state)
    _build_agenda_items(doc, state, designation_map)
    _build_action_summary(doc, state, designation_map)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        doc.save(tmp_path)
        upload_blob(tmp_path, blob_path)
        logger.info("doc_writer_node: uploaded → '%s'.", blob_path)
    except Exception as exc:
        logger.error("doc_writer_node: %s", exc)
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        return {**state, "error": f"doc_writer_node: {exc}"}
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    return {**state, "docx_path": blob_path, "docx_filename": filename}